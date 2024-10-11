import fitz  # PyMuPDF for PDF extraction  
from openai import AzureOpenAI  
from dotenv import load_dotenv  
import os  
import re  # For parsing structured output  
import json  # For JSON handling  
import pandas as pd  
import streamlit as st  
import docx  
from io import BytesIO  
  
# Load environment variables from .env file  
load_dotenv()  
  
# Set up Azure OpenAI API credentials from .env  
client = AzureOpenAI(  
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),  # Pull from environment  
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),  # Pull from environment  
    api_version=os.getenv("OPENAI_API_VERSION"),  # Pull from environment  
)  
  
# Function to extract text from a PDF file and format it as JSON  
def extract_text_from_pdf(pdf_path):  
    """Extract text from all pages of a PDF file and format it as JSON."""  
    try:  
        doc = fitz.open(pdf_path)  
        full_text = {"pages": []}  
        # Extract text from each page  
        for page_num in range(len(doc)):  
            page = doc.load_page(page_num)  
            page_text = page.get_text("text")  
            full_text["pages"].append({  
                "page_number": page_num + 1,  
                "text": page_text  
            })  
        doc.close()  
        return json.dumps(full_text)  # Convert the dictionary to a JSON string  
    except Exception as e:  
        print(f"Error extracting text from PDF: {e}")  
        return ""  
  
# Function to determine domain expertise needed  
def determine_domain_expertise(action_document_text):  
    """Analyze the action document to determine the required domain expertise."""  
    prompt = f"""  
    Analyze the following action document text and determine the domain expertise required to analyze this document:  
    {action_document_text}  
    Step 1: Identify the technical field and specific domain expertise required to understand the claims and references in the document.  
    Step 2: Return the domain expertise needed for the patent to be analysed in detail.  
    Step 3: Extract the main domain and subject matter needed and give the response as "You are a patent attorney specialising in ...." within 100 words. Display only from the "You.. " part.  
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": "You are a patent attorney analyzing the document to determine the required domain expertise."  
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    # Call OpenAI API for domain expertise determination  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
        domain_expertise = response.choices[0].message.content.strip()  
        return domain_expertise  
    except Exception as e:  
        print(f"Error during domain expertise determination: {e}")  
        return None  
  
# Function to check for conflicts using OpenAI API  
def check_for_conflicts(action_document_text, domain_expertise):  
    """  
    Analyzes the action document and extracts:  
    - Foundational claim  
    - Referenced documents  
    - Figures and technical text related to them  
    """  
    # Escape curly braces in the action_document_text  
    escaped_text = action_document_text.replace("{", "{{").replace("}", "}}")  
      
    prompt = f"""  
    Analyze the following action document text and extract the foundational claim:  
    {escaped_text}  
    Step 1: Extract the key claims from the document and name it as 'Key_claims'.  
    Step 2: From the 'Key_claims' extract the foundational claim and store it in a variable called "foundational_claim" (Note: method claims and system claims are not considered independent claims and only one claim can be the foundational claim).  
    Step 3: From the foundational claim, extract the information under U.S.C 102 and/or 103.  
    Step 4: Extract all referenced documents under U.S.C. 102 and/or 103 mentioned in the action document specified only in the "foundational_claim".  
    Step 5: For each referenced document, create a variable that stores the document name.  
    Step 6: If the foundational claim refers to the referenced documents, extract the entire technical content with its specified paragraph location and image reference. Map the claim with the conflicting document name.  
    Step 7: Do not extract any referenced document data that is not related to the foundational claim.   
    NOTE:Extract in English.  
    Step 8: Return the output as a JSON object with the following structure:  
    {{  
        "foundational_claim": "text",  
        "documents_referenced": ["doc1", "doc2", ...],  
        "figures": ["fig1", "fig2", ...],  
        "text": "detailed text"  
    }}  
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": f"You are a patent attorney with expertise in {domain_expertise} analyzing the document for foundational claims and conflicts."  
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    # Call OpenAI API for conflict checking  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
        # Extract the content and remove the triple backticks  
        content = response.choices[0].message.content.strip()  
          
        if content.startswith("```json"):  
            content = content[7:-3].strip()  
        elif content.startswith("```"):  
            content = content[3:-3].strip()  
          
        # Print the raw response for debugging  
        print(f"Raw response: {response.choices[0].message.content}")  
          
        # Parse the JSON content  
        return json.loads(content)  
      
    except json.JSONDecodeError as e:  
        print(f"JSON decoding error: {e}")  
        print(f"Raw response: {response.choices[0].message.content}")  
        return None  
    except Exception as e:  
        print(f"Error during conflict checking: {e}")  
        return None  
  
# Function to extract and analyze figure-related details  
def extract_figures_and_text(conflict_results, ref_document_text):  
    """  
    Extract figures and related technical text from the 'check_for_conflicts' function's output.  
    """  
    # Extract the 'figures' and 'text' sections from the JSON output  
    fig_details = conflict_results.get("figures", [])  
    text_details = conflict_results.get("text", "")  
      
    # Prepare a structured prompt for figure analysis  
    figure_analysis_prompt = f"""  
    Analyze the figures and technical text from the referenced document in relation to the foundational claim.  
    Instructions:  
    1. Identify Figures:  
        - For each figure referenced in the foundational claim, extract the following:  
            - **Figure Number and Title:** Provide the figure number and its title.  
            - **Technical Details:** Extract all technical details related to the figure as mentioned in the text. Ensure no technical detail is missed.  
            - **Importance:** Explain the importance of the figure in relation to the foundational claim. Describe how it supports, illustrates, or contradicts the claim.  
    2. Extract Text from Paragraphs:  
        - From the paragraphs cited in the foundational claim, extract the relevant text as in the document uploaded and store it in a separate variable.  
    3. Workflow for Cases with Images:  
        - If figures are present in the referenced document:  
            - Follow the steps outlined above to extract figure details and technical information.  
            - Ensure that any interpretations of the figures include specific references to the data or concepts depicted.  
    4. Workflow for Cases without Images:  
        - If no figures are present:  
            - Focus on extracting and analyzing the text from the referenced document.  
            - Identify and highlight key technical details and concepts that are essential to understanding the foundational claim.  
    Input Details:  
    Figures: {json.dumps(fig_details, indent=2)}  
    Text: {text_details}  
    Referenced Document Text: {ref_document_text}  
    Example Output Format:  
    {{  
        "figures_analysis": [  
            {{  
                "figure_number": "Figure 1",  
                "title": "Title of Figure 1",  
                "technical_details": "Detailed text",  
                "importance": "Explanation of importance"  
            }},  
            ...  
        ],  
        "extracted_paragraphs": [  
            "Paragraph text 1",  
            ...  
        ]  
    }}  
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": "You are a technical expert analyzing figures in a document."  
        },  
        {  
            "role": "user",  
            "content": figure_analysis_prompt,  
        },  
    ]  
      
    # Call OpenAI API for figure analysis  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
        analysis_output = response.choices[0].message.content.strip()  
          
        # Remove the triple backticks if they exist  
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
          
        # Print the raw response for debugging  
        print(f"Raw response: {response.choices[0].message.content}")  
          
        # Parse the JSON content  
        return json.loads(analysis_output)  
      
    except json.JSONDecodeError as e:  
        print(f"JSON decoding error: {e}")  
        print(f"Raw response: {response.choices[0].message.content}")  
        return None  
    except Exception as e:  
        print(f"Error during figure analysis: {e}")  
        return None  
  
# Function to extract details from the application as filed relating to the foundational claim  
def extract_details_from_filed_application(filed_application_text, foundational_claim):  
    """  
    Extract details from the filed application related to the foundational claim.  
    """  
    prompt = f"""  
    Analyze the following filed application text and extract details related to the foundational claim.  
    Filed Application Text: {filed_application_text}  
    Foundational Claim: {json.dumps(foundational_claim, indent=2)}  
    Instructions:  
    1. Identify and extract all technical details from the filed application that relate to the foundational claim.  
    2. Ensure that any extracted details include specific references to the paragraphs or sections in the filed application where they are found. NOTE:Extract in English.  
    3. Return the extracted details in the following JSON format:  
    {{  
        "foundational_claim_details": [  
            {{  
                "paragraph_number": "Paragraph 1",  
                "text": "Detailed text related to the foundational claim"  
            }},  
            ...  
        ]  
    }}  
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": "You are a patent attorney analyzing the filed application for details related to the foundational claim."  
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    # Call OpenAI API for extracting details from the filed application  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
        analysis_output = response.choices[0].message.content.strip()  
          
        # Remove the triple backticks if they exist  
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
          
        # Print the raw response for debugging  
        print(f"Raw response: {response.choices[0].message.content}")  
          
        # Parse the JSON content  
        return json.loads(analysis_output)  
      
    except json.JSONDecodeError as e:  
        print(f"JSON decoding error: {e}")  
        print(f"Raw response: {response.choices[0].message.content}")  
        return None  
    except Exception as e:  
        print(f"Error extracting details from filed application: {e}")  
        return None  
  
# Function to extract details from pending claims and modify the filed application details  
def extract_and_modify_filed_application(filed_application_details, pending_claims_text):  
    """  
    Extract details from the pending claims and modify the filed application details.  
    """  
    prompt = f"""  
    Analyze the following pending claims text and modify the filed application details accordingly.  
    Pending Claims Text: {pending_claims_text}  
    Filed Application Details: {json.dumps(filed_application_details, indent=2)}  
    Instructions:  
    1. Identify and extract all technical details from the pending claims that relate to the foundational claim.  
    2. Modify the filed application details based on the extracted details from the pending claims.  
    3. Ensure that any modifications include specific references to the paragraphs or sections in the pending claims where they are found.NOTE:Extract in English.  
    4. Return the modified filed application details in the following JSON format:  
    {{  
        "modified_filed_application_details": [  
            {{  
                "paragraph_number": "Paragraph 1",  
                "text": "Modified detailed text based on pending claims"  
            }},  
            ...  
        ]  
    }}  
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": "You are a patent attorney analyzing the pending claims to modify the filed application details."  
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    # Call OpenAI API for extracting and modifying filed application details  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
        analysis_output = response.choices[0].message.content.strip()  
          
        # Remove the triple backticks if they exist  
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
          
        # Print the raw response for debugging  
        print(f"Raw response: {response.choices[0].message.content}")  
          
        # Parse the JSON content  
        return json.loads(analysis_output)  
      
    except json.JSONDecodeError as e:  
        print(f"JSON decoding error: {e}")  
        print(f"Raw response: {response.choices[0].message.content}")  
        return None  
    except Exception as e:  
        print(f"Error extracting and modifying filed application details: {e}")  
        return None  
  
# Function to analyze the filed application based on the foundational claim, figure analysis, and application details  
def analyze_filed_application(extracted_details, foundational_claim, figure_analysis):  
    prompt = f"""  
    Analyze the filed application based on the foundational claim:  
    {json.dumps(foundational_claim, indent=2)}  
    and the figure analysis results:  
    {json.dumps(figure_analysis, indent=2)}  
    and the application as filed details:  
    {extracted_details}  
    Assess whether the examiner's rejection of the application under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified by comparing it with the cited references text.  
Instructions:  
Key Features of foundational claim  
Extract the key features of the foundational claimKey Features of Cited Reference  
Extract the key features of the cited reference  
Examiner’s Analysis  
Describe the examiner’s analysis and the basis for rejection.
Novelty Analysis  
Conduct a novelty analysis comparing the foundational claim with the cited reference.Non-obviousness Analysis  
Conduct a non-obviousness analysis comparing the foundational claim with the cited reference.Conclusion  
Provide a conclusion on whether the examiner’s rejection under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified.Potential Areas for Distinction Listed  
Identify potential areas for distinction in the foundational claim.
Proposed Amendments and Arguments  
Amendment to Foundational Claim for Each Key Feature Listed Separately with the New Features Highlighted:  
  Amendment [Number]: [Feature]     
  Original: "[Original feature description...]"     
  Proposed: "[Enhanced feature description with new details, specific materials, or configurations...]"  
  
Proposed Argument to the Examiner:  
Provide arguments for novelty and non-obviousness over the cited reference.Identify Limitations in the Current Claims  
Identify limitations in the current claims and propose specific language or structural changes that address those limitations.Propose New Arguments or Amendments  
Propose new arguments or amendments that distinguish the foundational claim from the cited prior art and strengthen the application. Ensure to include multiple amendments for thorough differentiation in depth.Maintain Original Intent  
Note:Ensure the amendments maintain the original intent of the claims while improving clarity and scope.
Note:Do the numbering in bullets and not in numbers.
Note: Wherever U.S.C 102 is mentioned, it should be printed as U.S.C 102 (Lack of Novelty), and wherever U.S.C 103 is mentioned, it should be printed as U.S.C 103 (Obviousness).  
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": (  
                "Adopt the persona of a Person Having Ordinary Skill in the Art (PHOSITA). "  
                "Analyze the filed application text and determine if the examiner is correct "  
                "in rejecting the application under either U.S.C 102 or U.S.C 103. "  
                "Cite instances from the application as filed to justify your stance."  
            )  
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
        analysis_output = response.choices[0].message.content.strip()  
          
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
          
        try:  
            return json.loads(analysis_output)  
        except json.JSONDecodeError:  
            return analysis_output  
    except Exception as e:  
        print(f"Error during filed application analysis: {e}")  
        return None  
  
def analyze_modified_application(cited_references_text, foundational_claim, figure_analysis, modified_application_details):  
    prompt = f"""  
    Analyze the modified application based on the foundational claim:{json.dumps(foundational_claim, indent=2)}and the figure analysis results:{json.dumps(figure_analysis, indent=2)}and the modified application details:{json.dumps(modified_application_details, indent=2)}and the cited references:{json.dumps(cited_references_text, indent=2)}  
Assess whether the examiner's rejection of the application under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified by comparing it with the cited references text.  
Instructions:  
Key Features of Pending Claim  
Extract the key features of the pending claimKey Features of Cited Reference  
Extract the key features of the cited reference  
Examiner’s Analysis  
Describe the examiner’s analysis and the basis for rejection.
Novelty Analysis  
Conduct a novelty analysis comparing the pending claim with the cited reference.
Non-obviousness Analysis  
Conduct a non-obviousness Analysis comparing the pending claim with the cited reference.
Conclusion  
Provide a conclusion on whether the examiner’s rejection under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified.Potential Areas for Distinction Listed  
Identify potential areas for distinction in the pending claim.Proposed Amendments and Arguments  
Amendment to Foundational Claim for Each Key Feature Listed Separately with the New Features Highlighted:  
  Amendment [Number]: [Feature]  
  Original: "[Original feature description...]"  
  Proposed: "[Enhanced feature description with new details, specific materials, or configurations...]"  
  
Proposed Argument to the Examiner:  
Provide arguments for novelty and non-obviousness over the cited reference.Identify Limitations in the Current Claims  
Identify limitations in the current claims and propose specific language or structural changes that address those limitations.Propose New Arguments or Amendments  
Propose new arguments or amendments that distinguish the foundational claim from the cited prior art and strengthen the application. Ensure to include multiple amendments for thorough differentiation in depth.Maintain Original Intent  
Note:Ensure the amendments maintain the original intent of the claims while improving clarity and scope.  
Note:Do the numbering in bullets and not in numbers.
Note: Wherever U.S.C 102 is mentioned, it should be printed as U.S.C 102 (Lack of Novelty), and wherever U.S.C 103 is mentioned, it should be printed as U.S.C 103 (Obviousness).  
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": (  
                "Adopt the persona of a Person Having Ordinary Skill in the Art (PHOSITA). "  
                "Analyze the modified application text and determine if the examiner is correct "  
                "in rejecting the application under either U.S.C 102 or U.S.C 103. "  
                "Cite instances from the modified application to justify your stance."  
            )  
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
        analysis_output = response.choices[0].message.content.strip()  
          
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
          
        try:  
            return json.loads(analysis_output)  
        except json.JSONDecodeError:  
            return analysis_output  
    except Exception as e:  
        print(f"Error during modified application analysis: {e}")  
        return None  
  
def save_analysis_to_word(analysis_output):
    if analysis_output is None or analysis_output.strip() == "":
        st.error("Analysis data is missing or empty.")
        return None

    # Create a new Word document
    doc = docx.Document()
    doc.add_heading('Filed Application Analysis Results', level=1)

    # Split the analysis output into lines
    lines = analysis_output.split('\n')
    for line in lines:
        line = line.strip()

        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(line, style='List Number')
        else:
            # Create a new paragraph for normal or mixed text (bold and non-bold)
            paragraph = doc.add_paragraph()

            # Use regex to find text between **...** for bold words
            # Split by bold sections while keeping bold markers for processing
            parts = re.split(r'(\*\*.*?\*\*)', line)
            
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    # This is the bold part, remove the '**' and set it as bold
                    bold_text = part[2:-2]
                    run = paragraph.add_run(bold_text)
                    run.bold = True
                else:
                    # This is regular text
                    paragraph.add_run(part)

    # Save the document to a BytesIO buffer instead of writing to disk
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer 
# Initialize session state variables  
if 'conflict_results' not in st.session_state:  
    st.session_state.conflict_results = None  
if 'foundational_claim' not in st.session_state:  
    st.session_state.foundational_claim = None  
if 'figure_analysis' not in st.session_state:  
    st.session_state.figure_analysis = None  
if 'filed_application_analysis' not in st.session_state:  
    st.session_state.filed_application_analysis = None  
if 'cited_documents' not in st.session_state:  
    st.session_state.cited_documents = None  
if 'pending_claims_analysis' not in st.session_state:  
    st.session_state.pending_claims_analysis = None  
if 'pending_claims_available' not in st.session_state:  
    st.session_state.pending_claims_available = "No"  # Default to "No"  
  
# Function to create aligned uploader and button  
def create_uploader_and_button(label_button, key):  
    col1, col2 = st.columns([4, 1])  # Adjust the column widths as needed  
    with col1:  
        uploaded_file = st.file_uploader("", type="pdf", key=key)  # Empty string for no label  
    with col2:  
        st.markdown("<br>", unsafe_allow_html=True)  # Add some space with HTML  
        button_clicked = st.button(label_button)  
    return uploaded_file, button_clicked  
  
# Step 1: Upload Examiner Document and Check Conflicts  
with st.expander("Step 1: Office Action", expanded=True):  
    st.write("### Upload the Examiner Document and Check for Conflicts")  
    uploaded_examiner_file, conflicts_clicked = create_uploader_and_button("Check for Conflicts", "examiner")  
  
    if conflicts_clicked:  
        if uploaded_examiner_file is not None:  
            with open("temp_examiner.pdf", "wb") as f:  
                f.write(uploaded_examiner_file.read())  
            extracted_examiner_text = extract_text_from_pdf("temp_examiner.pdf")  
            os.remove("temp_examiner.pdf")  
  
            # Domain expertise and conflict checking  
            domain_expertise = determine_domain_expertise(extracted_examiner_text)  
            if domain_expertise:  
                conflict_results_raw = check_for_conflicts(extracted_examiner_text, domain_expertise)  
                if conflict_results_raw:  
                    st.session_state.conflict_results = conflict_results_raw  
                    st.session_state.foundational_claim = conflict_results_raw.get("foundational_claim", "")  
                    st.session_state.cited_documents = conflict_results_raw.get("documents_referenced", [])  
                    st.success("Conflicts checked successfully!")  
                else:  
                    st.error("Failed to check for conflicts.")  
            else:  
                st.error("Failed to determine domain expertise.")  
        else:  
            st.warning("Please upload the examiner document first.")  
  
# Display Cited Documents Referenced after Step 1  
if st.session_state.get("cited_documents") is not None:  
    st.write("### Cited Documents Referenced:")  
    cited_docs_df = pd.DataFrame(st.session_state.cited_documents, columns=["Document Name"])  
    st.table(cited_docs_df)  
  
# Step 2: Upload Referenced Document and Analyze Figures  
if st.session_state.get("conflict_results") is not None:  # Ensure Step 1 was completed  
    with st.expander("Step 2: Referenced Documents", expanded=True):  
        st.write("### Upload the Referenced Document and Analyze Figures")  
        uploaded_ref_file, analyze_figures_clicked = create_uploader_and_button("Analyze Figures and Cited Text", "referenced")  
  
        if analyze_figures_clicked:  
            if uploaded_ref_file is not None:  
                with open("temp_referenced.pdf", "wb") as f:  
                    f.write(uploaded_ref_file.read())  
                extracted_ref_text = extract_text_from_pdf("temp_referenced.pdf")  
                os.remove("temp_referenced.pdf")  
  
                # Perform figure analysis  
                figure_analysis_results = extract_figures_and_text(st.session_state.conflict_results, extracted_ref_text)  
                if figure_analysis_results:  
                    st.session_state.figure_analysis = figure_analysis_results  
                    st.success("Figure analysis completed successfully!")  
                else:  
                    st.error("Failed to analyze figures and cited text.")  
            else:  
                st.warning("Please upload the referenced document first.")  
  
# Step 3: Upload Filed Application and Analyze  
if st.session_state.get("figure_analysis") is not None:  
    with st.expander("Step 3: Application as Filed", expanded=True):  
        st.write("### Upload the Filed Application and Analyze")  
        uploaded_filed_app, analyze_filed_app_clicked = create_uploader_and_button("Analyze Filed Application", "filed")  
  
        if analyze_filed_app_clicked:  
            if uploaded_filed_app is not None:  
                with open("temp_filed.pdf", "wb") as f:  
                    f.write(uploaded_filed_app.read())  
                extracted_filed_app_text = extract_text_from_pdf("temp_filed.pdf")  
                os.remove("temp_filed.pdf")  
  
                # Perform filed application analysis  
                filed_app_details = extract_details_from_filed_application(  
                    extracted_filed_app_text,  
                    st.session_state.foundational_claim  
                )  
                if filed_app_details:  
                    filed_app_details_json = json.dumps(filed_app_details, indent=2)  
                    st.session_state.filed_application_analysis = filed_app_details_json  
  
                    # Perform filed application analysis and generate report  
                    analysis_results = analyze_filed_application(  
                        filed_app_details_json,  
                        st.session_state.foundational_claim,  
                        st.session_state.figure_analysis  
                    )  
                    if analysis_results:  
                        st.session_state.filed_application_analysis = analysis_results  
                        st.success("Filed application analysis completed successfully!")  
                        docx_buffer = save_analysis_to_word(analysis_results)  
                        if docx_buffer:  
                            st.download_button(  
                                label="Download Analysis Results",  
                                data=docx_buffer,  
                                file_name="filed_application_analysis.docx",  
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
                                key="filed_application_download"  
                            )  
                    else:  
                        st.error("Failed to analyze the filed application.")  
                else:  
                    st.error("Failed to analyze the filed application.")  
            else:  
                st.warning("Please upload the filed application first.")  
  
# Step 4: Pending Claims  
if st.session_state.get("filed_application_analysis") is not None:  
    with st.expander("Step 4: Pending Claims", expanded=True):  
        st.write("### Do you have a Pending Claims Document to Analyze?")  
        st.session_state.pending_claims_available = st.radio("Select an option:", ("Yes", "No"), index=0 if st.session_state.pending_claims_available == "Yes" else 1)  
  
        if st.session_state.pending_claims_available == "Yes":  
            st.write("### Upload the Pending Claims Document and Analyze")  
            uploaded_pending_claims_file, analyze_pending_claims_clicked = create_uploader_and_button("Analyze Pending Claims", "pending_claims")  
  
            if analyze_pending_claims_clicked:  
                if uploaded_pending_claims_file is not None:  
                    with open("temp_pending_claims.pdf", "wb") as f:  
                        f.write(uploaded_pending_claims_file.read())  
                    extracted_pending_claims_text = extract_text_from_pdf("temp_pending_claims.pdf")  
                    os.remove("temp_pending_claims.pdf")  
  
                    # Perform pending claims analysis  
                    pending_claims_analysis_results = analyze_modified_application(  
                        extracted_pending_claims_text,  
                        st.session_state.foundational_claim,  
                        st.session_state.figure_analysis,  
                        st.session_state.filed_application_analysis  
                    )  
                    if pending_claims_analysis_results:  
                        st.session_state.pending_claims_analysis = pending_claims_analysis_results  
                        st.success("Pending claims analysis completed successfully!")  
                        st.json(pending_claims_analysis_results)  
  
                        # Generate report for download  
                        docx_buffer = save_analysis_to_word(pending_claims_analysis_results)  
                        if docx_buffer:  
                            st.download_button(  
                                label="Download Analysis Results",  
                                data=docx_buffer,  
                                file_name="pending_claims_analysis.docx",  
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
                                key="pending_claims_download"  
                            )  
                    else:  
                        st.error("Failed to analyze the pending claims.")  
                else:  
                    st.warning("Please upload the pending claims document first.")  
  
# Option to download results if there are no pending claims  
if st.session_state.get("filed_application_analysis") and st.session_state.pending_claims_analysis is None:  
    docx_buffer = save_analysis_to_word(st.session_state.filed_application_analysis)  
    if docx_buffer:  
        st.download_button(  
            label="Download Analysis Results",  
            data=docx_buffer,  
            file_name="filed_application_analysis.docx",  
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
            key="filed_application_final_download"  
        )  
